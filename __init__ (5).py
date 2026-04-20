from __future__ import annotations

from pathlib import Path
import zipfile

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


def text_to_docx(text: str, output_path: Path, title: str | None = None) -> Path:
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for line in text.splitlines() or [text]:
        doc.add_paragraph(line)
    doc.save(output_path)
    return output_path


def text_to_pdf(text: str, output_path: Path, title: str | None = None) -> Path:
    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    margin = 50
    y = height - margin
    if title:
        c.setFont("Helvetica-Bold", 14)
        c.drawString(margin, y, title[:120])
        y -= 24
    c.setFont("Helvetica", 10)
    max_width = width - (margin * 2)

    for raw_line in text.splitlines() or [text]:
        line = raw_line.strip() or " "
        words = line.split()
        current = ""
        for word in words or [""]:
            trial = (current + " " + word).strip()
            if stringWidth(trial, "Helvetica", 10) <= max_width:
                current = trial
            else:
                c.drawString(margin, y, current)
                y -= 14
                current = word
                if y <= margin:
                    c.showPage()
                    c.setFont("Helvetica", 10)
                    y = height - margin
        c.drawString(margin, y, current)
        y -= 14
        if y <= margin:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = height - margin
    c.save()
    return output_path


def zip_folder(folder: Path, zip_path: Path) -> Path:
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for file_path in folder.rglob("*"):
            if file_path.is_file():
                zf.write(file_path, arcname=file_path.relative_to(folder))
    return zip_path

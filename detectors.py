from pathlib import Path

from docx import Document

from app.anonymizer.service import AnonymisationService
from app.processors.base import ProcessedArtifact
from app.utils.exporters import text_to_pdf


class DocxProcessor:
    def process(self, input_path: Path, output_dir: Path, service: AnonymisationService) -> ProcessedArtifact:
        document = Document(str(input_path))
        preview_parts: list[str] = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                result = service.anonymize_text(paragraph.text)
                paragraph.text = result.anonymized_text
                preview_parts.append(result.anonymized_text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        result = service.anonymize_text(cell.text)
                        cell.text = result.anonymized_text
                        preview_parts.append(result.anonymized_text)

        for section in document.sections:
            for paragraph in section.header.paragraphs:
                if paragraph.text.strip():
                    paragraph.text = service.anonymize_text(paragraph.text).anonymized_text
            for paragraph in section.footer.paragraphs:
                if paragraph.text.strip():
                    paragraph.text = service.anonymize_text(paragraph.text).anonymized_text

        stem = input_path.stem + "_anonymised"
        docx_path = output_dir / f"{stem}.docx"
        pdf_path = output_dir / f"{stem}.pdf"
        document.save(docx_path)
        preview_text = "\n".join(preview_parts)
        text_to_pdf(preview_text, pdf_path, title=input_path.name)
        return ProcessedArtifact(input_path=input_path, output_files=[docx_path, pdf_path], preview_text=preview_text[:3000])
